import os
import json
import re
import fitz
import io
import docx
import jwt
from jwt import PyJWKClient
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends, Security
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv

from qdrant_client import QdrantClient
from langchain_huggingface import HuggingFaceEndpointEmbeddings
from langchain_qdrant import QdrantVectorStore
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

load_dotenv()

SARVAM_API_KEY = os.getenv("SARVAM_API_KEY")
QDRANT_URL = os.getenv("QDRANT_URL")
QDRANT_API_KEY = os.getenv("QDRANT_API_KEY")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
HUGGINGFACEHUB_API_TOKEN = os.getenv("HUGGINGFACEHUB_API_TOKEN")

# ── FIXED: Added rstrip('/') to prevent double-slash bugs ──
raw_sub_url = os.getenv("NEXT_PUBLIC_SUPABASE_URL") or os.getenv("SUPABASE_URL") or ""
SUPABASE_URL = raw_sub_url.rstrip("/")
FRONTEND_URL = os.getenv("FRONTEND_URL", "http://localhost:3000") 

app = FastAPI(title="Vidhi-Vichara Enterprise Core")

# ── SECURE CORS POLICY (UPDATED WITH ALL PRODUCTION DOMAINS) ──
allowed_origins = [
    FRONTEND_URL,
    "http://localhost:3000",
    "https://vidhivichara.in",
    "https://www.vidhivichara.in",
    "https://vidhivichara.netlify.app"
]

app.add_middleware(
    CORSMiddleware, 
    allow_origins=allowed_origins, 
    allow_credentials=True,
    allow_methods=["*"], 
    allow_headers=["*"] 
)

security = HTTPBearer()

# ── LOCAL JWKS CLIENT SETUP (OFFICIAL SUPABASE PATH) ──
JWKS_URL = f"{SUPABASE_URL}/auth/v1/.well-known/jwks.json"
jwks_client = PyJWKClient(JWKS_URL)

# ── BULLETPROOF SUPABASE VERIFICATION (HANDLES ECC ROTATION) ──
def verify_token(credentials: HTTPAuthorizationCredentials = Security(security)):
    raw_token = credentials.credentials
    try:
        # Dynamically grab the public key that matches the ECC token header
        signing_key = jwks_client.get_signing_key_from_jwt(raw_token)
        
        # Decode using the modern ES256 algorithm
        payload = jwt.decode(
            raw_token, 
            signing_key.key, 
            algorithms=["ES256", "RS256", "HS256"], 
            audience="authenticated"
        )
        
        # Extract user email and assign RBAC roles
        email = payload.get("email", "unknown_user")
        role = "Administrator" if email == "admin.iitjodpur.vidhivichara2026@gmail.com" else "Verified User"
        
        return {"role": role, "email": email}
        
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token has expired. Please log in again.")
    except Exception as e:
        print(f"Auth Verification Error: {e}")
        raise HTTPException(status_code=401, detail="Invalid or expired authentication token.")

# ── AI ENGINE CONFIGURATION ──
q_client = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY, timeout=60.0)

# Replaced local PyTorch model with cloud endpoint to prevent 512MB RAM crash on Render
embeddings = HuggingFaceEndpointEmbeddings(
    model="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
    huggingfacehub_api_token=HUGGINGFACEHUB_API_TOKEN
)

llm = ChatGroq(api_key=GROQ_API_KEY, model_name="llama-3.3-70b-versatile", temperature=0.0).bind(response_format={"type": "json_object"})
vectorstore = QdrantVectorStore(client=q_client, collection_name="indian_laws_master", embedding=embeddings)

vvai_prompt = """\
You are an expert Indian Constitutional Law AI. Evaluate the vires, scope, and constitutional validity of the 'Proposed Rule' based on Schedule VII (Union, State, Concurrent Lists).

USER'S SELECTED JURISDICTION FROM UI: **{jurisdiction}**

STEP 1: INDEPENDENT CLASSIFICATION
Determine if this is inherently a Central or State Act. If State, name the specific State.
STEP 2: MISMATCH DETECTION
Compare findings to User Jurisdiction ({jurisdiction}). Flag cross-over errors as 'Jurisdictional Overreach'.

OUTPUT STRICTLY AS A VALID JSON OBJECT.
{{
  "detected_parent_act": "Extract the specific parent Act name from the text (e.g., 'Information Technology Act, 2000') or 'General Law'",
  "detected_jurisdiction": "Central" | "State" | "Unknown",
  "detected_state": "Name of State or 'N/A'",
  "jurisdiction_mismatch": true | false,
  "vvai_score": <float>, 
  "band": "Green | Amber | Red | Critical | Unknown",
  "dimensions": {{"D1_Vires": 0.0, "D2_Definitional": 0.0, "D3_Semantic": 0.0, "D4_Scope": 0.0, "D5_Procedural": 0.0, "D6_Sanction": 0.0, "D7_Purposive": 0.0}},
  "deviation_type": "T1|T2|T3|T4|T5|Jurisdictional Overreach|None", 
  "severity": "S1|S2|S3|S4|Critical|None",
  "violating_quote": "Extract exact violating text.",
  "explanation": "Assessment summary.", 
  "suggested_fix": "Remediation guide."
}}
Context: {context}"""

prompt = ChatPromptTemplate.from_messages([("system", vvai_prompt), ("human", "Proposed Rule: {input}")])

remediate_prompt_template = """\
You are an expert Parliamentary Draftsman specializing in Indian Constitutional Law. 
You are given a 'Flagged Clause' that has been marked as a Jurisdictional Overreach violation because a State entity tried to legislate on a Central subject, or vice versa.

Your task is to completely rewrite this clause to bring it into absolute harmony with the Indian Constitution, while preserving the user's core intent.

OUTPUT STRICTLY AS A VALID JSON OBJECT:
{{
  "original_clause": "{flagged_clause}",
  "compliant_draft": "The completely rewritten, constitutionally valid draft text.",
  "drafting_notes": "A technical legal explanation of the drafting modifications made to bypass the vires violation."
}}"""

remediate_prompt = ChatPromptTemplate.from_messages([("system", remediate_prompt_template)])

def format_docs(docs): 
    return "\n\n".join(f"[Act: {d.metadata.get('act_name', d.metadata.get('source', 'Unknown Act'))}]\n{d.page_content}" for d in docs)

# ── BULLETPROOF EXTRACTOR (.pdf, .txt, .docx with Tables) ──
def extract_text(file_bytes: bytes, filename: str) -> str:
    text = ""
    filename = filename.lower()
    try:
        if filename.endswith(".pdf"):
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in doc: text += page.get_text()
        elif filename.endswith(".txt"): 
            text = file_bytes.decode('utf-8')
        elif filename.endswith(".docx"):
            doc = docx.Document(io.BytesIO(file_bytes))
            text_chunks = [para.text for para in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_chunks.append(cell.text)
            text = "\n".join(text_chunks)
        else:
            raise ValueError("Unsupported file format. Please upload .pdf, .txt, or .docx")
    except Exception as e:
        print(f"Extraction Error: {e}")
        raise ValueError("Failed to parse document content. Ensure the file is not corrupted.")
        
    return text

# ── CORE API ENDPOINTS ──

@app.post("/api/audit")
async def run_audit(jurisdiction: str = Form(...), target_act: str = Form(...), file: UploadFile = File(...), user: dict = Depends(verify_token)):
    try:
        file_bytes = await file.read()
        
        if len(file_bytes) > 5_000_000:
            raise HTTPException(status_code=400, detail="File too large. Please upload a document under 5MB.")
            
        rule_text = extract_text(file_bytes, file.filename)
        if not rule_text.strip(): raise ValueError("Document is empty.")

        detected_act = target_act
        retriever = vectorstore.as_retriever(search_kwargs={"k": 3})
        retrieved_docs = retriever.invoke(rule_text[:1500])
        
        if target_act == "AUTO_DETECT":
            act_pattern = r"((?:[A-Z][a-zA-Z]*\s+)+(?:(?:of|to|and|for|the|in|on|at)\s+)?(?:[A-Z][a-zA-Z]*\s+)*Act,?\s*\d{4})"
            act_match = re.search(act_pattern, rule_text)
            
            if act_match:
                detected_act = act_match.group(1).strip()
            else:
                if retrieved_docs:
                    meta = retrieved_docs[0].metadata
                    raw_name = meta.get("act_name") or meta.get("source") or meta.get("title") or meta.get("file_name") or "General Law"
                    detected_act = "General Law" if raw_name == "General Law" else os.path.basename(raw_name).replace(".pdf", "").replace(".txt", "").replace("_", " ").title()
                else:
                    detected_act = "General Law"

        context_str = format_docs(retrieved_docs)
        chain = prompt | llm | StrOutputParser()
        raw_result = chain.invoke({"jurisdiction": jurisdiction.title(), "context": context_str, "input": rule_text[:5000]})
        final_report = json.loads(raw_result)
        
        ai_detected_act = final_report.get("detected_parent_act", "")
        if detected_act == "General Law" and ai_detected_act and ai_detected_act not in ["", "Unknown", "General Law"]:
            detected_act = ai_detected_act

        final_report["detected_act"] = detected_act
        return final_report
        
    except ValueError as ve:
        raise HTTPException(status_code=400, detail=str(ve))
    except HTTPException:
        raise
    except Exception as e:
        print(f"CRITICAL BACKEND ERROR: {e}")
        raise HTTPException(status_code=500, detail="Internal server error processing the evaluation.")

@app.post("/api/remediate")
async def remediate_clause(flagged_clause: str = Form(...), user: dict = Depends(verify_token)):
    try:
        chain = remediate_prompt | llm | StrOutputParser()
        raw_result = chain.invoke({"flagged_clause": flagged_clause})
        return json.loads(raw_result)
    except Exception as e:
        print(f"REMEDIATION ERROR: {e}")
        raise HTTPException(status_code=500, detail="Drafting Engine encountered an internal error.")

@app.get("/api/admin/benchmark")
async def run_system_benchmark(user: dict = Depends(verify_token)):
    if user["role"] != "Administrator": raise HTTPException(status_code=403, detail="Forbidden Access")
    
    golden_dataset = [
        {"input": "The State Government of Maharashtra hereby rules that all international data traffic crossing local servers shall be subject to state encryption taxation.", "expected_mismatch": True},
        {"input": "The Ministry of Finance mandates corporations to file national tax balance details under Central protocols.", "expected_mismatch": False},
        {"input": "The Legislative Assembly of Tamil Nadu bans commercial aircraft movements over municipal limits during night hours.", "expected_mismatch": True}
    ]
    
    passed_tests = 0
    results_log = []
    
    for idx, test in enumerate(golden_dataset):
        chain = prompt | llm | StrOutputParser()
        raw_result = chain.invoke({"jurisdiction": "Central", "context": "System Benchmark Baseline", "input": test["input"]})
        report = json.loads(raw_result)
        
        is_correct = report.get("jurisdiction_mismatch") == test["expected_mismatch"]
        if is_correct: passed_tests += 1
        
        results_log.append({"test_id": idx + 1, "verdict": "PASSED" if is_correct else "FAILED", "detected_jurisdiction": report.get("detected_jurisdiction")})
        
    accuracy_percentage = (passed_tests / len(golden_dataset)) * 100
    return {"status": "Benchmark Evaluation Complete", "total_test_cases": len(golden_dataset), "accuracy_score": f"{accuracy_percentage:.1f}%", "detailed_metrics": results_log}

@app.post("/api/admin/ingest")
async def ingest_act(act_name: str = Form(...), file: UploadFile = File(...), user: dict = Depends(verify_token)):
    if user["role"] != "Administrator": raise HTTPException(status_code=403, detail="Forbidden")
    try:
        file_bytes = await file.read()
        text = extract_text(file_bytes, file.filename)
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
        chunks = splitter.split_text(text)
        
        docs = [Document(page_content=c, metadata={"act_name": act_name}) for c in chunks] 
        vectorstore.add_documents(docs)
        return {"status": "success", "message": f"Successfully ingested {len(docs)} segments."}
    except Exception as e:
        print(f"INGESTION ERROR: {e}")
        raise HTTPException(status_code=500, detail="Failed to ingest document into vector database.")